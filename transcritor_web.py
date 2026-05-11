# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import tempfile
import time
import zipfile
from datetime import timedelta
from pathlib import Path

import streamlit as st


APP_TITLE = "Transcritor Web EN-PT"
VERSION = "1.0"

BASE_DIR = Path(__file__).resolve().parent
MODELS_DIR = BASE_DIR / "models"
OUTPUT_DIR = BASE_DIR / "saida_web"
MODELS_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


st.set_page_config(
    page_title=APP_TITLE,
    page_icon="🎧",
    layout="wide",
    initial_sidebar_state="collapsed",
)


def aplicar_estilo():
    st.markdown(
        """
        <style>
        html, body, [data-testid="stAppViewContainer"] { overflow-y: auto !important; }
        .block-container {
            max-width: 1180px;
            padding: 1.4rem 1rem 5rem 1rem;
        }
        h1 { font-size: 1.65rem !important; margin-bottom: .2rem !important; }
        h2, h3 { letter-spacing: 0 !important; }
        [data-testid="stHeader"] { background: rgba(255,255,255,.96); }
        div[data-testid="stMetric"] {
            background: #f8fafc;
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            padding: .55rem .65rem;
        }
        div[data-testid="stAlert"] { border-radius: 8px; }
        .hero {
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            padding: .9rem 1rem;
            background: #ffffff;
            margin: .6rem 0 1rem 0;
        }
        .hero strong { display: block; font-size: 1rem; margin-bottom: .2rem; }
        .hero span { color: #64748b; }
        .result-card {
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            padding: .85rem;
            background: #ffffff;
            margin: .7rem 0;
        }
        @media (max-width: 640px) {
            .block-container { padding: 1rem .55rem 5rem .55rem; }
            h1 { font-size: 1.25rem !important; line-height: 1.2 !important; }
            h2 { font-size: 1.08rem !important; }
            h3 { font-size: 1rem !important; }
            p, div, span { font-size: .92rem; }
            div[data-testid="column"] { min-width: 0 !important; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def format_timestamp(seconds: float) -> str:
    if seconds < 0:
        seconds = 0
    td = timedelta(seconds=float(seconds))
    total = int(td.total_seconds())
    millis = int((float(seconds) - int(float(seconds))) * 1000)
    hours = total // 3600
    minutes = (total % 3600) // 60
    secs = total % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def safe_name(filename: str) -> str:
    stem = Path(filename).stem
    name = "".join(ch if ch.isalnum() or ch in " ._-()" else "_" for ch in stem).strip()
    return name or "audio"


def write_txt(path: Path, title: str, lines: list[str]) -> None:
    content = [title, "=" * len(title), ""]
    content.extend(lines)
    path.write_text("\n".join(content).strip() + "\n", encoding="utf-8")


def write_srt(path: Path, segments: list[dict], lines: list[str] | None = None) -> None:
    blocks = []
    for i, seg in enumerate(segments, 1):
        text = lines[i - 1] if lines else seg["text"]
        blocks.append(
            f"{i}\n"
            f"{format_timestamp(seg['start'])} --> {format_timestamp(seg['end'])}\n"
            f"{text.strip()}\n"
        )
    path.write_text("\n".join(blocks).strip() + "\n", encoding="utf-8")


def group_by_pause(segments: list[dict], lines: list[str], pause_seconds: float = 0.8) -> list[str]:
    paragraphs: list[str] = []
    current: list[str] = []
    previous_end: float | None = None

    for seg, line in zip(segments, lines):
        text = (line or "").strip()
        if not text:
            continue

        gap = 0.0 if previous_end is None else float(seg["start"]) - previous_end
        if current and gap >= pause_seconds:
            paragraphs.append(" ".join(current).strip())
            current = []

        current.append(text)
        previous_end = float(seg["end"])

    if current:
        paragraphs.append(" ".join(current).strip())

    return paragraphs


def write_docx(
    path: Path,
    source_name: str,
    english_paragraphs: list[str],
    portuguese_paragraphs: list[str] | None = None,
) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    doc.add_heading("Transcricao e traducao", level=1)
    doc.add_paragraph(f"Arquivo de origem: {source_name}")

    if portuguese_paragraphs is not None:
        doc.add_heading("Traducao em portugues", level=2)
        for paragraph in portuguese_paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    doc.add_heading("Transcricao em ingles", level=2)
    for paragraph in english_paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    doc.save(path)


def load_argos_translation():
    try:
        import argostranslate.translate
    except Exception as exc:
        return None, f"Argos Translate nao instalado: {exc}"

    try:
        installed = argostranslate.translate.get_installed_languages()
        source = next((lang for lang in installed if lang.code == "en"), None)
        target = next((lang for lang in installed if lang.code == "pt"), None)
        if not source or not target:
            return None, "Pacote Argos en->pt ainda nao instalado."
        translation = source.get_translation(target)
        if not translation:
            return None, "Traducao Argos en->pt nao encontrada."
        return translation, "Argos en->pt pronto."
    except Exception as exc:
        return None, f"Falha ao carregar Argos en->pt: {exc}"


def instalar_argos_en_pt() -> str:
    import argostranslate.package

    argostranslate.package.update_package_index()
    packages = argostranslate.package.get_available_packages()
    package = next((p for p in packages if p.from_code == "en" and p.to_code == "pt"), None)
    if package is None:
        raise RuntimeError("Pacote Argos en->pt nao encontrado.")
    path = package.download()
    argostranslate.package.install_from_path(path)
    return "Pacote Argos en->pt instalado."


def translate_lines_argos(lines: list[str]) -> list[str]:
    translator, status = load_argos_translation()
    if translator is None:
        raise RuntimeError(status)
    translated: list[str] = []
    progress = st.progress(0, text="Traduzindo para portugues...")
    total = max(len(lines), 1)
    for i, line in enumerate(lines, 1):
        clean = line.strip()
        translated.append(translator.translate(clean) if clean else "")
        if i == 1 or i == total or i % 5 == 0:
            progress.progress(i / total, text=f"Traduzindo linha {i}/{total}...")
    progress.empty()
    return translated


@st.cache_resource(show_spinner=False)
def carregar_modelo(model_name: str, device: str, compute_type: str, local_only: bool):
    from faster_whisper import WhisperModel

    return WhisperModel(
        model_name,
        device=device,
        compute_type=compute_type,
        download_root=str(MODELS_DIR),
        local_files_only=local_only,
    )


def transcrever(input_path: Path, model_name: str, device: str, compute_type: str, local_only: bool):
    model = carregar_modelo(model_name, device, compute_type, local_only)
    segments_iter, info = model.transcribe(
        str(input_path),
        language="en",
        task="transcribe",
        vad_filter=True,
        beam_size=5,
    )

    segments: list[dict] = []
    english_lines: list[str] = []
    status = st.empty()
    for seg in segments_iter:
        text = seg.text.strip()
        segments.append({"start": float(seg.start), "end": float(seg.end), "text": text})
        english_lines.append(text)
        if len(segments) == 1 or len(segments) % 10 == 0:
            status.info(f"Transcrevendo trecho {len(segments)}...")
    status.empty()

    if not segments:
        raise RuntimeError("Nenhum texto foi encontrado no audio/video.")

    return segments, english_lines, info


def salvar_upload(uploaded_file) -> Path:
    suffix = Path(uploaded_file.name).suffix or ".mp3"
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    temp.write(uploaded_file.getbuffer())
    temp.close()
    return Path(temp.name)


def gerar_arquivos(
    source_name: str,
    segments: list[dict],
    english_lines: list[str],
    portuguese_lines: list[str] | None,
    pause_seconds: float,
) -> tuple[list[Path], str, str]:
    base = safe_name(source_name)
    stamp = time.strftime("%Y%m%d_%H%M%S")
    out_dir = OUTPUT_DIR / f"{base}_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    english_paragraphs = group_by_pause(segments, english_lines, pause_seconds)
    portuguese_paragraphs = (
        group_by_pause(segments, portuguese_lines, pause_seconds) if portuguese_lines is not None else None
    )

    created: list[Path] = []
    english_txt = out_dir / f"{base}_transcricao_en.txt"
    write_txt(english_txt, "Transcricao em ingles", english_paragraphs)
    created.append(english_txt)

    if portuguese_paragraphs is not None:
        pt_txt = out_dir / f"{base}_traducao_pt.txt"
        write_txt(pt_txt, "Traducao em portugues", portuguese_paragraphs)
        created.append(pt_txt)

    en_srt = out_dir / f"{base}_legenda_en.srt"
    write_srt(en_srt, segments)
    created.append(en_srt)

    if portuguese_lines is not None:
        pt_srt = out_dir / f"{base}_legenda_pt.srt"
        write_srt(pt_srt, segments, portuguese_lines)
        created.append(pt_srt)

    docx_path = out_dir / f"{base}_transcricao_traducao.docx"
    write_docx(docx_path, source_name, english_paragraphs, portuguese_paragraphs)
    created.append(docx_path)

    zip_path = out_dir / f"{base}_resultado.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in created:
            zf.write(path, arcname=path.name)

    preview_pt = "\n\n".join(portuguese_paragraphs or [])
    preview_en = "\n\n".join(english_paragraphs)
    return [zip_path] + created, preview_pt, preview_en


def limpar_texto_tamanho(nome: str) -> str:
    return re.sub(r"\s+", " ", nome).strip()


def main():
    aplicar_estilo()

    st.title(f"{APP_TITLE} {VERSION}")
    st.caption("Transcricao de audio/video em ingles, traducao para portugues e geracao de legenda.")
    st.markdown(
        """
        <div class="hero">
            <strong>Primeira versao online do seu transcritor</strong>
            <span>Envie um arquivo, transcreva, traduza e baixe TXT, SRT, Word e ZIP.</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.header("Configuracao")
        model_name = st.selectbox("Modelo Whisper", ["tiny", "base", "small", "medium"], index=1)
        device = st.selectbox("Dispositivo", ["cpu", "cuda"], index=0)
        compute_type = st.selectbox("Precisao", ["int8", "float16", "float32"], index=0)
        local_only = st.checkbox("Usar somente modelos ja baixados", value=False)
        pause_seconds = st.slider("Pausa para paragrafo", 0.5, 5.0, 0.8, 0.1)
        traduzir = st.checkbox("Traduzir para portugues", value=True)

        translator, argos_status = load_argos_translation()
        st.caption(argos_status)
        if translator is None and st.button("Instalar traducao EN-PT"):
            with st.spinner("Instalando pacote de traducao..."):
                try:
                    st.success(instalar_argos_en_pt())
                    st.rerun()
                except Exception as exc:
                    st.error(f"Nao foi possivel instalar: {exc}")

    uploaded = st.file_uploader(
        "Envie um audio ou video em ingles",
        type=["mp3", "wav", "m4a", "aac", "flac", "ogg", "mp4", "mkv", "mov", "avi", "webm"],
    )

    if uploaded is None:
        st.info("Escolha um arquivo para comecar.")
        return

    c1, c2, c3 = st.columns(3)
    c1.metric("Arquivo", limpar_texto_tamanho(uploaded.name)[:32])
    c2.metric("Tamanho", f"{uploaded.size / (1024 * 1024):.1f} MB")
    c3.metric("Modelo", model_name)

    if not st.button("Transcrever agora", type="primary", use_container_width=True):
        return

    temp_path: Path | None = None
    try:
        temp_path = salvar_upload(uploaded)
        with st.spinner("Carregando modelo e transcrevendo..."):
            segments, english_lines, info = transcrever(temp_path, model_name, device, compute_type, local_only)

        portuguese_lines: list[str] | None = None
        if traduzir:
            portuguese_lines = translate_lines_argos(english_lines)

        created, preview_pt, preview_en = gerar_arquivos(
            uploaded.name,
            segments,
            english_lines,
            portuguese_lines,
            pause_seconds,
        )

        st.success("Concluido. Arquivos prontos para download.")
        m1, m2, m3 = st.columns(3)
        m1.metric("Trechos", len(segments))
        m2.metric("Idioma", getattr(info, "language", "en"))
        m3.metric("Arquivos", len(created) - 1)

        zip_path = created[0]
        st.download_button(
            "Baixar tudo em ZIP",
            data=zip_path.read_bytes(),
            file_name=zip_path.name,
            mime="application/zip",
            use_container_width=True,
        )

        tab1, tab2, tab3 = st.tabs(["Portugues", "Ingles", "Arquivos"])
        with tab1:
            if preview_pt:
                st.text_area("Previa da traducao", preview_pt, height=320)
            else:
                st.info("Traducao nao gerada nesta execucao.")
        with tab2:
            st.text_area("Previa da transcricao", preview_en, height=320)
        with tab3:
            for path in created[1:]:
                mime = "application/octet-stream"
                if path.suffix == ".txt":
                    mime = "text/plain"
                elif path.suffix == ".srt":
                    mime = "text/plain"
                elif path.suffix == ".docx":
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                st.download_button(path.name, path.read_bytes(), file_name=path.name, mime=mime)

    except Exception as exc:
        st.error("Nao foi possivel concluir a transcricao.")
        st.exception(exc)
    finally:
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except OSError:
                pass


if __name__ == "__main__":
    main()
