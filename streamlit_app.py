import re
import base64
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from urllib.parse import urlparse

import streamlit as st
from docx import Document
from faster_whisper import WhisperModel


APP_TITLE = "Transcritor EN-PT"
VERSION = "v2.0"
OUTPUT_DIR = Path("outputs")
COOKIES_FILE = Path("cookies.txt")
SECRET_COOKIES_KEY = "YOUTUBE_COOKIES"
SECRET_COOKIES_B64_KEY = "YOUTUBE_COOKIES_B64"

SUPPORTED_EXTENSIONS = {
    ".mp3",
    ".wav",
    ".m4a",
    ".aac",
    ".flac",
    ".ogg",
    ".opus",
    ".mp4",
    ".mkv",
    ".mov",
    ".avi",
    ".webm",
}

YOUTUBE_DOMAINS = {
    "youtube.com",
    "www.youtube.com",
    "m.youtube.com",
    "youtu.be",
    "music.youtube.com",
}


@dataclass
class Segmento:
    start: float
    end: float
    text: str


def aplicar_estilo() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="🎙️", layout="wide")
    st.markdown(
        """
        <style>
            .block-container {
                padding-top: 2rem;
                max-width: 1120px;
            }
            .status-box {
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding: .85rem 1rem;
                background: #f9fafb;
                color: #374151;
                margin: .75rem 0 1rem;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


@st.cache_resource(show_spinner=False)
def carregar_modelo(
    model_name: str,
    device: str,
    compute_type: str,
    local_only: bool,
) -> WhisperModel:
    return WhisperModel(
        model_name,
        device=device,
        compute_type=compute_type,
        local_files_only=local_only,
    )


def limpar_texto(nome: str) -> str:
    return re.sub(r"\s+", " ", nome).strip()


def nome_seguro(nome: str) -> str:
    stem = Path(nome).stem
    stem = re.sub(r"[^\w\-. ]+", "", stem, flags=re.UNICODE)
    stem = re.sub(r"\s+", "_", stem).strip("._-")
    return stem or "transcricao"


def url_valida(url: str) -> bool:
    parsed = urlparse(url.strip())
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def eh_youtube(url: str) -> bool:
    domain = urlparse(url.strip()).netloc.lower()
    return domain in YOUTUBE_DOMAINS or domain.endswith(".youtube.com")


def salvar_upload(uploaded_file) -> Path:
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Formato nao suportado: {suffix}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        return Path(tmp.name)


def baixar_link_direto(url: str, temp_dir: Path) -> tuple[Path, str]:
    try:
        import requests
    except ImportError as exc:
        raise RuntimeError("Instale 'requests' no requirements.txt.") from exc

    parsed = urlparse(url.strip())
    suffix = Path(parsed.path).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("O link nao aponta diretamente para um arquivo de audio/video suportado.")

    original_name = Path(parsed.path).name or f"midia{suffix}"
    file_path = temp_dir / f"{nome_seguro(original_name)}{suffix}"

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept": "*/*",
    }

    try:
        with requests.get(
            url,
            headers=headers,
            stream=True,
            timeout=60,
            allow_redirects=True,
        ) as response:
            response.raise_for_status()
            content_type = response.headers.get("content-type", "").lower()

            if "text/html" in content_type:
                raise RuntimeError("O link retornou uma pagina HTML, nao um arquivo de midia.")

            with file_path.open("wb") as file:
                for chunk in response.iter_content(chunk_size=1024 * 1024):
                    if chunk:
                        file.write(chunk)
    except Exception as exc:
        raise RuntimeError(
            "Nao foi possivel baixar esse link direto. Confira se ele termina com "
            ".mp3, .mp4, .wav, .m4a, .aac, .flac, .ogg, .opus ou .webm."
        ) from exc

    if not file_path.exists() or file_path.stat().st_size == 0:
        raise RuntimeError("O arquivo baixado esta vazio.")

    return file_path, original_name


def explicar_erro_ytdlp(error_text: str, youtube: bool) -> str:
    texto = error_text.lower()

    if "sign in to confirm" in texto or "not a bot" in texto or "captcha" in texto:
        return (
            "O YouTube pediu confirmacao de acesso. Em servidores como Streamlit Cloud, "
            "isso costuma acontecer. Exporte cookies do seu navegador para um arquivo "
            "chamado cookies.txt e coloque na raiz do projeto."
        )

    if "private video" in texto or "members-only" in texto:
        return "Esse video parece ser privado, restrito ou exclusivo para membros."

    if "age-restricted" in texto or "age restricted" in texto:
        return "Esse video tem restricao de idade. Use cookies.txt de uma conta que consiga assistir."

    if "this video is unavailable" in texto or "video unavailable" in texto:
        return "Esse video esta indisponivel para o servidor, pode ser por regiao, remocao ou privacidade."

    if "ffmpeg" in texto:
        return "O servidor precisa do FFmpeg instalado. No Streamlit Cloud, adicione um packages.txt com ffmpeg."

    if youtube:
        return (
            "O YouTube recusou o download neste ambiente. Tente cookies.txt, use upload manual "
            "do arquivo, ou cole um link direto para .mp3/.mp4."
        )

    return "O site recusou ou bloqueou o download do link."


def cookies_disponiveis() -> bool:
    return (
        COOKIES_FILE.exists()
        or bool(st.secrets.get(SECRET_COOKIES_KEY, "").strip())
        or bool(st.secrets.get(SECRET_COOKIES_B64_KEY, "").strip())
    )


def preparar_cookiefile(temp_dir: Path) -> Path | None:
    secret_cookies = st.secrets.get(SECRET_COOKIES_KEY, "").strip()
    secret_cookies_b64 = st.secrets.get(SECRET_COOKIES_B64_KEY, "").strip()

    if secret_cookies:
        cookie_path = temp_dir / "youtube_cookies.txt"
        cookie_path.write_text(secret_cookies, encoding="utf-8")
        return cookie_path

    if secret_cookies_b64:
        try:
            decoded = base64.b64decode(secret_cookies_b64).decode("utf-8")
        except Exception as exc:
            raise RuntimeError("YOUTUBE_COOKIES_B64 esta invalido nos Secrets do Streamlit.") from exc

        cookie_path = temp_dir / "youtube_cookies.txt"
        cookie_path.write_text(decoded, encoding="utf-8")
        return cookie_path

    if COOKIES_FILE.exists():
        return COOKIES_FILE

    return None


def baixar_com_ytdlp(url: str, temp_dir: Path) -> tuple[Path, str]:
    try:
        from yt_dlp import YoutubeDL
        from yt_dlp.utils import DownloadError
    except ImportError as exc:
        raise RuntimeError("Instale 'yt-dlp' no requirements.txt.") from exc

    ydl_opts = {
        "format": "bestaudio[ext=m4a]/bestaudio[ext=mp3]/bestaudio/best",
        "outtmpl": str(temp_dir / "%(title).200B.%(ext)s"),
        "noplaylist": True,
        "quiet": True,
        "no_warnings": True,
        "retries": 5,
        "fragment_retries": 5,
        "socket_timeout": 30,
        "concurrent_fragment_downloads": 1,
        "extractor_args": {"youtube": {"player_client": ["android", "web"]}},
        "http_headers": {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0.0.0 Safari/537.36"
            ),
            "Accept": "*/*",
            "Accept-Language": "pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7",
        },
    }

    cookiefile = preparar_cookiefile(temp_dir)

    if cookiefile is not None:
        ydl_opts["cookiefile"] = str(cookiefile)

    try:
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
    except DownloadError as exc:
        error_text = str(exc)
        explanation = explicar_erro_ytdlp(error_text, eh_youtube(url))
        raise RuntimeError(f"{explanation}\n\nDetalhe tecnico do yt-dlp: {error_text}") from exc

    arquivos = [path for path in temp_dir.iterdir() if path.is_file()]

    if not arquivos:
        raise RuntimeError("Nenhum arquivo foi baixado a partir do link.")

    arquivo_baixado = max(arquivos, key=lambda path: path.stat().st_size)
    source_name = limpar_texto(info.get("title") or arquivo_baixado.name)

    return arquivo_baixado, source_name


def baixar_midia_link(url: str) -> tuple[Path, str, Path]:
    if not url_valida(url):
        raise ValueError("Informe um link valido com http:// ou https://.")

    temp_dir = Path(tempfile.mkdtemp(prefix="transcritor_link_"))

    try:
        parsed = urlparse(url.strip())
        suffix = Path(parsed.path).suffix.lower()

        if suffix in SUPPORTED_EXTENSIONS:
            file_path, source_name = baixar_link_direto(url, temp_dir)
        else:
            file_path, source_name = baixar_com_ytdlp(url, temp_dir)

        return file_path, source_name, temp_dir
    except Exception:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise


def transcrever(
    audio_path: Path,
    model_name: str,
    device: str,
    compute_type: str,
    local_only: bool,
) -> tuple[list[Segmento], list[str], object]:
    model = carregar_modelo(model_name, device, compute_type, local_only)
    raw_segments, info = model.transcribe(
        str(audio_path),
        language="en",
        vad_filter=True,
        beam_size=5,
    )

    segmentos: list[Segmento] = []
    linhas: list[str] = []

    for segment in raw_segments:
        texto = segment.text.strip()
        if not texto:
            continue

        segmentos.append(Segmento(float(segment.start), float(segment.end), texto))
        linhas.append(texto)

    if not segmentos:
        raise RuntimeError("Nenhum trecho de fala foi detectado no arquivo.")

    return segmentos, linhas, info


@st.cache_resource(show_spinner=False)
def garantir_argos_en_pt() -> str:
    try:
        import argostranslate.package
        import argostranslate.translate
    except ImportError as exc:
        raise RuntimeError("Instale 'argostranslate' no requirements.txt.") from exc

    installed_languages = argostranslate.translate.get_installed_languages()
    from_lang = next((lang for lang in installed_languages if lang.code == "en"), None)
    to_lang = next((lang for lang in installed_languages if lang.code == "pt"), None)

    if from_lang is not None and to_lang is not None:
        try:
            from_lang.get_translation(to_lang)
            return "Traducao EN -> PT disponivel."
        except Exception:
            pass

    argostranslate.package.update_package_index()
    available_packages = argostranslate.package.get_available_packages()
    package = next(
        (pkg for pkg in available_packages if pkg.from_code == "en" and pkg.to_code == "pt"),
        None,
    )

    if package is None:
        raise RuntimeError("Pacote EN -> PT nao encontrado no indice do Argos.")

    package_path = package.download()
    argostranslate.package.install_from_path(package_path)

    return "Pacote de traducao EN -> PT instalado."


def load_argos_translation():
    try:
        import argostranslate.translate
    except ImportError:
        return None, "Argos Translate nao esta instalado."

    installed_languages = argostranslate.translate.get_installed_languages()
    from_lang = next((lang for lang in installed_languages if lang.code == "en"), None)
    to_lang = next((lang for lang in installed_languages if lang.code == "pt"), None)

    if from_lang is None or to_lang is None:
        return None, "Pacote de traducao EN -> PT ainda nao instalado."

    try:
        translator = from_lang.get_translation(to_lang)
    except Exception:
        return None, "Pacote EN -> PT encontrado, mas nao pode ser carregado."

    return translator, "Traducao EN -> PT pronta."


def translate_lines_argos(lines: Iterable[str]) -> list[str]:
    translator, status = load_argos_translation()

    if translator is None:
        raise RuntimeError(status)

    return [translator.translate(line).strip() for line in lines]


def agrupar_paragrafos(
    segments: list[Segmento],
    lines: list[str],
    pause_seconds: float,
) -> list[str]:
    if not segments or not lines:
        return []

    paragraphs: list[str] = []
    current: list[str] = [lines[0]]

    for index in range(1, min(len(segments), len(lines))):
        pause = segments[index].start - segments[index - 1].end
        if pause >= pause_seconds:
            paragraphs.append(" ".join(current).strip())
            current = [lines[index]]
        else:
            current.append(lines[index])

    if current:
        paragraphs.append(" ".join(current).strip())

    return paragraphs


def format_srt_time(seconds: float) -> str:
    milliseconds = int(round(seconds * 1000))
    hours = milliseconds // 3_600_000
    milliseconds %= 3_600_000
    minutes = milliseconds // 60_000
    milliseconds %= 60_000
    secs = milliseconds // 1000
    millis = milliseconds % 1000
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def write_srt(
    path: Path,
    segments: list[Segmento],
    lines: list[str] | None = None,
) -> None:
    with path.open("w", encoding="utf-8") as file:
        for index, segment in enumerate(segments, start=1):
            text = lines[index - 1] if lines and index - 1 < len(lines) else segment.text
            file.write(f"{index}\n")
            file.write(f"{format_srt_time(segment.start)} --> {format_srt_time(segment.end)}\n")
            file.write(f"{text.strip()}\n\n")


def write_txt(path: Path, title: str, paragraphs: list[str]) -> None:
    with path.open("w", encoding="utf-8") as file:
        file.write(f"{title}\n")
        file.write("=" * len(title))
        file.write("\n\n")
        file.write("\n\n".join(paragraphs))


def write_docx(
    path: Path,
    source_name: str,
    english_paragraphs: list[str],
    portuguese_paragraphs: list[str] | None,
) -> None:
    doc = Document()
    doc.add_heading("Transcricao e traducao", level=1)
    doc.add_paragraph(f"Arquivo original: {source_name}")
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    if portuguese_paragraphs:
        doc.add_heading("Traducao em portugues", level=2)
        for paragraph in portuguese_paragraphs:
            doc.add_paragraph(paragraph)

    doc.add_heading("Transcricao em ingles", level=2)
    for paragraph in english_paragraphs:
        doc.add_paragraph(paragraph)

    doc.save(path)


def gerar_arquivos(
    source_name: str,
    segments: list[Segmento],
    english_lines: list[str],
    portuguese_lines: list[str] | None,
    pause_seconds: float,
) -> tuple[list[Path], str, str]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    base = nome_seguro(source_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = OUTPUT_DIR / f"{base}_{timestamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    english_paragraphs = agrupar_paragrafos(segments, english_lines, pause_seconds)
    portuguese_paragraphs = (
        agrupar_paragrafos(segments, portuguese_lines, pause_seconds)
        if portuguese_lines is not None
        else None
    )

    created: list[Path] = []

    english_txt = out_dir / f"{base}_transcricao_en.txt"
    write_txt(english_txt, "Transcricao em ingles", english_paragraphs)
    created.append(english_txt)

    if portuguese_paragraphs is not None:
        pt_txt = out_dir / f"{base}_traducao_pt.txt"
        write_txt(pt_txt, "Traducao em portugues", portuguese_paragraphs)
        created.append(pt_txt)

    en_srt = out_dir / f"{base}_legenda_en.srt"
    write_srt(en_srt, segments)
    created.append(en_srt)

    if portuguese_lines is not None:
        pt_srt = out_dir / f"{base}_legenda_pt.srt"
        write_srt(pt_srt, segments, portuguese_lines)
        created.append(pt_srt)

    docx_path = out_dir / f"{base}_transcricao_traducao.docx"
    write_docx(docx_path, source_name, english_paragraphs, portuguese_paragraphs)
    created.append(docx_path)

    zip_path = out_dir / f"{base}_resultado.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in created:
            zf.write(path, arcname=path.name)

    preview_pt = "\n\n".join(portuguese_paragraphs or [])
    preview_en = "\n\n".join(english_paragraphs)

    return [zip_path] + created, preview_pt, preview_en


def mostrar_ajuda_youtube() -> None:
    st.markdown(
        """
        <div class="status-box">
            <strong>Sobre links do YouTube</strong><br>
            O app tenta usar yt-dlp. Se o YouTube bloquear o servidor, use upload manual
            ou configure YOUTUBE_COOKIES_B64 nos secrets do Streamlit. Links diretos para
            arquivos .mp3 ou .mp4 costumam funcionar sem cookies.
        </div>
        """,
        unsafe_allow_html=True,
    )


def main() -> None:
    aplicar_estilo()

    st.title(f"{APP_TITLE} {VERSION}")
    st.caption("Transcreve audio/video em ingles, traduz para portugues e gera TXT, SRT, Word e ZIP.")

    with st.sidebar:
        st.header("Configuracao")

        model_name = st.selectbox("Modelo Whisper", ["tiny", "base", "small", "medium"], index=1)
        device = st.selectbox("Dispositivo", ["cpu", "cuda"], index=0)
        compute_type = st.selectbox("Precisao", ["int8", "float16", "float32"], index=0)
        local_only = st.checkbox("Usar somente modelos ja baixados", value=False)
        pause_seconds = st.slider("Pausa para paragrafo", 0.5, 5.0, 0.8, 0.1)
        traduzir = st.checkbox("Traduzir para portugues", value=True)
        instalar_auto = st.checkbox("Instalar traducao EN-PT automaticamente", value=True)
        mostrar_diagnostico = st.checkbox("Mostrar diagnostico tecnico de links", value=True)

        if traduzir and instalar_auto:
            try:
                with st.spinner("Verificando traducao EN -> PT..."):
                    st.caption(garantir_argos_en_pt())
            except Exception as exc:
                st.warning(f"Traducao automatica indisponivel: {exc}")

        translator, argos_status = load_argos_translation()
        st.caption(argos_status)

        if translator is None and st.button("Instalar traducao EN-PT"):
            with st.spinner("Instalando pacote de traducao..."):
                garantir_argos_en_pt.clear()
                st.success(garantir_argos_en_pt())
                st.rerun()

    modo_entrada = st.radio(
        "Escolha a origem do audio/video",
        ["Enviar arquivo", "Usar link"],
        horizontal=True,
    )

    uploaded = None
    link_midia = ""

    if modo_entrada == "Enviar arquivo":
        uploaded = st.file_uploader(
            "Envie um audio ou video em ingles",
            type=[
                "mp3",
                "wav",
                "m4a",
                "aac",
                "flac",
                "ogg",
                "opus",
                "mp4",
                "mkv",
                "mov",
                "avi",
                "webm",
            ],
        )

        if uploaded is None:
            st.info("Escolha um arquivo para comecar.")
            return

        c1, c2, c3 = st.columns(3)
        c1.metric("Arquivo", limpar_texto(uploaded.name)[:32])
        c2.metric("Tamanho", f"{uploaded.size / (1024 * 1024):.1f} MB")
        c3.metric("Modelo", model_name)

    else:
        link_midia = st.text_input(
            "Cole o link do audio ou video",
            placeholder="https://www.youtube.com/watch?v=... ou https://site.com/audio.mp3",
        )
        mostrar_ajuda_youtube()

        if link_midia.strip() and eh_youtube(link_midia):
            if cookies_disponiveis():
                st.success("Cookies encontrados. O app vai tentar usar esses cookies no YouTube.")
            else:
                st.warning(
                    "YouTube detectado. Sem YOUTUBE_COOKIES_B64 nos secrets, "
                    "o Streamlit Cloud pode ser bloqueado."
                )

        if not link_midia.strip():
            st.info("Cole um link para comecar.")
            return

        c1, c2, c3 = st.columns(3)
        c1.metric("Origem", "YouTube" if eh_youtube(link_midia) else "Link")
        c2.metric("Cookies", "Sim" if cookies_disponiveis() else "Nao")
        c3.metric("Modelo", model_name)

    if not st.button("Transcrever agora", type="primary", use_container_width=True):
        return

    temp_path: Path | None = None
    temp_dir_link: Path | None = None
    source_name = ""

    try:
        if modo_entrada == "Enviar arquivo":
            temp_path = salvar_upload(uploaded)
            source_name = uploaded.name
        else:
            with st.spinner("Baixando midia do link..."):
                temp_path, source_name, temp_dir_link = baixar_midia_link(link_midia)

        with st.spinner("Carregando modelo e transcrevendo..."):
            segments, english_lines, _info = transcrever(
                temp_path,
                model_name,
                device,
                compute_type,
                local_only,
            )

        portuguese_lines: list[str] | None = None
        traducao_falhou = False

        if traduzir:
            with st.spinner("Traduzindo para portugues..."):
                try:
                    portuguese_lines = translate_lines_argos(english_lines)
                except Exception as exc:
                    traducao_falhou = True
                    st.warning(
                        "A transcricao foi concluida, mas a traducao nao foi realizada. "
                        f"Motivo: {exc}"
                    )

        with st.spinner("Gerando arquivos..."):
            files, preview_pt, preview_en = gerar_arquivos(
                source_name,
                segments,
                english_lines,
                portuguese_lines,
                pause_seconds,
            )

        tab1, tab2 = st.tabs(["Portugues", "Ingles"])

        with tab1:
            if portuguese_lines and not traducao_falhou:
                st.text_area("Traducao", preview_pt, height=320)
            else:
                st.info("Traducao indisponivel.")

        with tab2:
            st.text_area("Transcricao original", preview_en, height=320)

        st.success("Processamento concluido. Baixe os arquivos abaixo.")

        for file_path in files:
            with file_path.open("rb") as file:
                st.download_button(
                    label=f"Baixar {file_path.name}",
                    data=file,
                    file_name=file_path.name,
                    mime="application/octet-stream",
                )

    except Exception as exc:
        message = str(exc)

        if not mostrar_diagnostico and "Detalhe tecnico do yt-dlp:" in message:
            message = message.split("Detalhe tecnico do yt-dlp:", maxsplit=1)[0].strip()

        st.error(message)

    finally:
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except Exception:
                pass

        if temp_dir_link and temp_dir_link.exists():
            shutil.rmtree(temp_dir_link, ignore_errors=True)


if __name__ == "__main__":
    main()
